#####빅데이터플랫폼아키텍처 설계
###김용현

#문항1
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

filename = '빅데이터플랫폼 구축 kyh.docx'
document = Document()
document.add_paragraph('파이썬을 활용한 Word문서 작성')
document.save(filename)

first_paragraph = document.paragraphs[0]
first_run = first_paragraph.runs[0]
first_run.bold = True
first_run.font.size = Pt(24)
first_run.font.name = '맑은 고딕'
first_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
document.save(filename)

#문항2
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
filename = '빅데이터플랫폼 구축 kyh.docx'
document = Document(filename)
document.add_paragraph('김용현')
document.save(filename)

first_paragraph = document.paragraphs[1]
first_run = first_paragraph.runs[0]
first_run.underline = True
first_run.font.size = Pt(20)
first_run.font.name = '굴림'
first_run._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림')
document.save(filename)

#문항3
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
filename = '빅데이터플랫폼 구축 kyh.docx'
document = Document(filename)
new_paragraph = document.add_paragraph()
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
logo_run = new_paragraph.add_run()
logo_run.add_picture('Tjoeun_logo2.jpg')
logo_run.add_break(WD_BREAK.LINE)
caption_run = new_paragraph.add_run('[더조은 컴퓨터 아카데미]')
document.save(filename)

#문항4
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
filename = '빅데이터플랫폼 구축 kyh.docx'
document = Document(filename)
document.add_paragraph('훈련중입니다.')
document.save(filename)

first_paragraph = document.paragraphs[3]
first_run = first_paragraph.runs[0]
first_run.italic = True
first_run.font.size = Pt(40)
first_run.font.name = '맑은 고딕'
first_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
document.save(filename)